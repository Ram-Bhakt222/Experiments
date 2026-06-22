"""
pipelines/docs/docx_extract.py — DOCX -> Markdown + per-table CSV.

Preserves headings, paragraphs, lists, and tables. Output is a single .md file
plus tables/*.csv.
"""
from __future__ import annotations

import csv
from pathlib import Path
from typing import Any, Dict

from pipelines._common import first_input, job_output_dir
from queue_db import update_job


def _para_to_md(p) -> str:
    style = (p.style.name or "").lower()
    text = p.text or ""
    if style.startswith("heading"):
        try:
            level = int(style.split()[-1])
        except Exception:
            level = 1
        return "#" * max(1, min(6, level)) + " " + text
    if style.startswith("list") or text.startswith(("• ", "- ", "* ")):
        return "- " + text.lstrip("•-* ").strip()
    return text


def run(job: Dict[str, Any]) -> None:
    from docx import Document

    src = first_input(job)
    out = job_output_dir(job)
    update_job(job["id"], progress="reading docx")

    doc = Document(str(src))
    md_lines = []
    for p in doc.paragraphs:
        md_lines.append(_para_to_md(p))

    tables_dir = out / "tables"
    for ti, table in enumerate(doc.tables, start=1):
        tables_dir.mkdir(parents=True, exist_ok=True)
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        out_csv = tables_dir / f"table_{ti:03d}.csv"
        with out_csv.open("w", newline="", encoding="utf-8") as f:
            csv.writer(f).writerows(rows)

    md_path = out / f"{src.stem}.md"
    md_path.write_text("\n\n".join(l for l in md_lines if l.strip()), encoding="utf-8")

    update_job(
        job["id"],
        output_path=str(md_path),
        progress=f"done ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)",
    )
